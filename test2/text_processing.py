#!/usr/bin/env python3
"""
text_processing.py
---------------------
Extract text from TXT, DOCX, or PDF, split into paragraphs, and ensure each paragraph
is at most N words by splitting after sentence-ending periods.

Usage:
    python text_processing.py INPUT_FILE [--max-words 100] [--out OUTFILE] [--json JSONFILE]

Examples:
    python text_processing.py mydoc.pdf --max-words 120 --out output.txt
    python text_processing.py notes.docx --json output.json

Dependencies (install only what you need):
    pip install python-docx pdfplumber
"""
import argparse
import json
import re
import sys
from pathlib import Path
from typing import List

# --------------------- Text extraction helpers ---------------------

def extract_text_from_txt(path: Path) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text_from_docx(path: Path) -> str:
    try:
        from docx import Document  # python-docx
    except Exception as e:
        raise RuntimeError(
            "python-docx is required for .docx files. Install with: pip install python-docx"
        ) from e
    doc = Document(str(path))
    # Join paragraphs with double newlines so we can keep paragraph boundaries
    return "\n\n".join(p.text for p in doc.paragraphs)

def extract_text_from_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except Exception as e:
        raise RuntimeError(
            "pdfplumber is required for .pdf files. Install with: pip install pdfplumber"
        ) from e
    texts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            texts.append(t)
    raw = "\n".join(texts)
    # Heuristic cleanup to turn line-wraps into spaces but keep paragraph breaks
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"(?<!\n)\n(?!\n)", " ", raw)   # single newlines -> space
    raw = re.sub(r"\n{3,}", "\n\n", raw)         # squash >2 newlines to 2
    return raw

def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return extract_text_from_txt(path)
    elif suffix == ".docx":
        return extract_text_from_docx(path)
    elif suffix == ".pdf":
        return extract_text_from_pdf(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use .txt, .docx, or .pdf")

# --------------------- Paragraph & sentence logic ---------------------

_WORD_RE = re.compile(r"\b[\w'-]+\b", flags=re.UNICODE)

def word_count(text: str) -> int:
    return len(_WORD_RE.findall(text))

# Split only on periods that likely end a sentence (your requirement).
# We avoid splitting on a few common abbreviations.
_ABBREV = r"(?:Mr|Mrs|Ms|Dr|Prof|Sr|Jr|vs|etc|e\.g|i\.e)"
_SENTENCE_SPLIT_RE = re.compile(
    rf"""
    (.*?)
    (?:
        (?<!\b{_ABBREV})
        \.
        (?=\s|$)
    )
    """,
    re.VERBOSE | re.DOTALL | re.IGNORECASE,
)

def split_into_sentences(text: str) -> List[str]:
    if "." not in text:
        return [text.strip()] if text.strip() else []
    sentences = []
    start = 0
    for m in _SENTENCE_SPLIT_RE.finditer(text):
        chunk = (m.group(0) or "").strip()
        if chunk:
            sentences.append(chunk)
        start = m.end()
    tail = text[start:].strip()
    if tail:
        sentences.append(tail)
    return sentences

def normalize_paragraphs(raw: str) -> List[str]:
    s = raw.replace("\r\n", "\n").replace("\r", "\n")
    s = "\n".join(line.strip() for line in s.split("\n"))
    s = re.sub(r"\n{2,}", "\n\n", s)
    parts = [p.strip() for p in s.split("\n\n")]

    # Merge very short “heading-like” paras into the next block
    merged = []
    for p in parts:
        if not p:
            continue
        if merged and word_count(merged[-1]) < 15:
            merged[-1] = (merged[-1] + " " + p).strip()
        else:
            merged.append(p)
    if not merged and raw.strip():
        return [raw.strip()]
    return merged

def enforce_max_words(paragraph: str, max_words: int) -> List[str]:
    if word_count(paragraph) <= max_words:
        return [paragraph.strip()]

    sentences = split_into_sentences(paragraph)
    if len(sentences) == 1:
        # No useful periods—fallback: chunk by words
        words = paragraph.split()
        return [" ".join(words[i:i+max_words]) for i in range(0, len(words), max_words)]

    chunks, current = [], []
    current_len = 0
    for s in sentences:
        s_len = word_count(s)
        if s_len > max_words:
            # If a single sentence exceeds limit, hard-split by words
            if current:
                chunks.append(" ".join(current).strip())
                current, current_len = [], 0
            words = s.split()
            chunks.extend(" ".join(words[i:i+max_words]) for i in range(0, len(words), max_words))
            continue
        if current_len + s_len <= max_words:
            current.append(s)
            current_len += s_len
        else:
            chunks.append(" ".join(current).strip())
            current, current_len = [s], s_len
    if current:
        chunks.append(" ".join(current).strip())
    return chunks

def process_text(raw: str, max_words: int) -> List[str]:
    paragraphs = normalize_paragraphs(raw)
    out: List[str] = []
    for p in paragraphs:
        out.extend(enforce_max_words(p, max_words))
    return [q.strip() for q in out if q and q.strip()]

# --------------------- CLI ---------------------

def main():
    ap = argparse.ArgumentParser(description="Extract and paragraphize a text-like document.")
    ap.add_argument("input_file", type=str, help="Path to .txt, .docx, or .pdf file")
    ap.add_argument("--max-words", type=int, default=100, help="Max words per paragraph (default: 100)")
    ap.add_argument("--out", type=str, default=None, help="Write the resulting .txt here (default: <input>.paragraphs.txt)")
    ap.add_argument("--json", type=str, default=None, help="Also write a JSON array of paragraphs to this path")
    args = ap.parse_args()

    inp = Path(args.input_file)
    if not inp.exists():
        print(f"File not found: {inp}", file=sys.stderr)
        sys.exit(1)

    try:
        raw = extract_text(inp)
    except Exception as e:
        print(f"Error extracting text: {e}", file=sys.stderr)
        sys.exit(2)

    paragraphs = process_text(raw, args.max_words)

    out_path = Path(args.out) if args.out else inp.with_suffix(".paragraphs.txt")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n\n".join(paragraphs) + ("\n" if paragraphs else ""), encoding="utf-8")

    if args.json:
        jp = Path(args.json)
        jp.parent.mkdir(parents=True, exist_ok=True)
        jp.write_text(json.dumps(paragraphs, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Wrote paragraphs to: {out_path}")
    if args.json:
        print(f"Wrote JSON to: {jp}")

if __name__ == "__main__":
    main()